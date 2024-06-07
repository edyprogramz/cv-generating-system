from django.shortcuts import render, redirect, get_object_or_404, HttpResponse
from resumes.models import Template as TemplateModel, TemplateCategory
from django.template import Template, Context
from django.contrib.auth.decorators import login_required
from weasyprint import HTML
# from django.template.loader import render_to_string
from client.models import Client, ProfessionalSummary, Education, Experience, Skill, Hobby, Reference, Language, Certification
# pdf 
# from xhtml2pdf import pisa
# from django.template.loader import get_template
from django.http import HttpResponse
# DOCX
from docx import Document
from bs4 import BeautifulSoup

# Create your views here.
def index(request):
    return render(request, "home.html", {})

def generate_pdf(request, client_id, template_id):
    client = get_object_or_404(Client, id=client_id)

    if not client.subscription_active():
        return redirect('payments:paypal_view', client_id=client_id, template_id=template_id)

    # Fetch the template data
    template_model = get_object_or_404(TemplateModel, id=template_id)

    # Fetch related data
    experiences = Experience.objects.filter(client=client)
    education_list = Education.objects.filter(client=client)
    skills = Skill.objects.filter(client=client)
    professional_summary = get_object_or_404(ProfessionalSummary, client=client)
    languages = Language.objects.filter(client=client)
    references = Reference.objects.filter(client=client)
    hobbies = Hobby.objects.filter(client=client)
    certifications = Certification.objects.filter(client=client)

    # Read the template content from the file
    with open(template_model.template_file.path, 'r', encoding='utf-8') as file:
        template_content = file.read()

    # Create a Template object
    template = Template(template_content)

    # Create a context with the fetched data
    context = Context({
        'client': client,
        'experiences': experiences,
        'education_list': education_list,
        'skills': skills,
        'professional_summary': professional_summary,
        'languages':languages,
        'references':references,
        'hobbies':hobbies,
        'certifications':certifications
    })

    # Render the template with the context
    html = template.render(context)

    # Create a PDF from the rendered HTML
    pdf_file = HTML(string=html).write_pdf()

    # Create an HTTP response object with PDF content type
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{client.first_name}_{client.last_name}_resume.pdf"'

    return response

def generate_docx(request, client_id, template_id):
    client = get_object_or_404(Client, id=client_id)
    template = get_object_or_404(TemplateModel, id=template_id)

    # Fetch client's data
    experiences = Experience.objects.filter(client=client)
    education_list = Education.objects.filter(client=client)
    skills = Skill.objects.filter(client=client)
    professional_summary = get_object_or_404(ProfessionalSummary, client=client)
    # added
    languages = Language.objects.filter(client=client)
    references = Reference.objects.filter(client=client)
    hobbies = Hobby.objects.filter(client=client)

    # Create a new Document
    document = Document()

    # Add client information
    document.add_heading(f'{client.first_name} {client.last_name}', 0)
    document.add_paragraph(f'Email: {client.email}')
    document.add_paragraph(f'Phone: {client.phone_number}')
    document.add_paragraph(f'Location: {client.city}, {client.country}')

    # Add professional summary
    document.add_heading('Professional Summary', level=1)
    document.add_paragraph(professional_summary.professional_summary)

    # Add experiences
    document.add_heading('Experience', level=1)
    for experience in experiences:
        document.add_heading(experience.job_title, level=2)
        document.add_paragraph(f'Company: {experience.company}')
        document.add_paragraph(f'Location: {experience.city}, {experience.country}')
        document.add_paragraph(f'Start Date: {experience.start_month} {experience.start_year}')
        document.add_paragraph(f'End Date: {experience.end_month} {experience.end_year}')

    # Add education
    document.add_heading('Education', level=1)
    for education in education_list:
        document.add_heading(education.institution_name, level=2)
        document.add_paragraph(f'Location: {education.institution_location}')
        document.add_paragraph(f'Program: {education.program_pursued}')
        document.add_paragraph(f'Field of Study: {education.field_of_study}')
        document.add_paragraph(f'Graduated: {"Yes" if education.graduated else "No"}')
        if education.graduation_month and education.graduation_year:
            document.add_paragraph(f'Graduation Date: {education.graduation_month} {education.graduation_year}')

    # Add skills
    document.add_heading('Skills', level=1)
    for skill in skills:
        document.add_paragraph(skill.name)

    # Prepare response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{client.first_name}_{client.last_name}_resume.docx"'
    document.save(response)

    return response

